from docx import Document
import os

def split_docx(filepath):
    doc = Document(filepath)
    output_dir = os.path.join("output", os.path.splitext(os.path.basename(filepath))[0])
    os.makedirs(output_dir, exist_ok=True)
    current_doc = None
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            if current_doc:
                current_doc.save(os.path.join(output_dir, f"{title}.docx"))
            title = para.text.strip().replace('/', '_').replace('\\', '_')
            current_doc = Document()
            current_doc.add_paragraph(para.text, style=para.style.name)
        elif current_doc:
            current_doc.add_paragraph(para.text, style=para.style.name)
    if current_doc:
        current_doc.save(os.path.join(output_dir, f"{title}.docx"))
    return output_dir